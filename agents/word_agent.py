"""
Word Agent — Microsoft Word automation via COM and python-docx fallback.

Capabilities:
- Create professional reports from scratch
- Open and edit existing documents
- Insert headings, paragraphs, tables, images, charts
- Rewrite content using LLM
- Apply professional formatting
- Save as DOCX or PDF
"""
from __future__ import annotations

import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from agents.base_agent import BaseAgent
from models.schemas import AgentType, RiskLevel
from utils.helpers import ensure_dir, timestamped_filename
from utils.logger import get_logger

logger = get_logger("agents.word", "word")

_COM_AVAILABLE = False
if sys.platform == "win32":
    try:
        import win32com.client as win32
        import pythoncom
        _COM_AVAILABLE = True
    except ImportError:
        logger.warning("pywin32 not available — using python-docx fallback")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not available")


class WordAgent(BaseAgent):
    """Automates Microsoft Word document creation and editing."""

    def __init__(self):
        self._word_app = None       # COM Word Application
        self._document = None       # COM or python-docx Document
        self._doc_path: Optional[str] = None
        self._use_com = _COM_AVAILABLE and os.getenv("DEMO_MODE", "false").lower() != "true"
        super().__init__(AgentType.WORD)

    def _register_tools(self):
        self.register_tool("word.create_document",  self.create_document,  "Create a new Word document")
        self.register_tool("word.open_document",    self.open_document,    "Open an existing Word document", ["path"])
        self.register_tool("word.insert_heading",   self.insert_heading,   "Insert a heading", ["text", "level"])
        self.register_tool("word.insert_paragraph", self.insert_paragraph, "Insert a paragraph", ["text"])
        self.register_tool("word.insert_table",     self.insert_table,     "Insert a data table", ["data"])
        self.register_tool("word.insert_image",     self.insert_image,     "Insert an image", ["image_path"])
        self.register_tool("word.insert_page_break",self.insert_page_break,"Insert a page break")
        self.register_tool("word.set_header",       self.set_header,       "Set document header text", ["text"])
        self.register_tool("word.set_footer",       self.set_footer,       "Set document footer text", ["text"])
        self.register_tool("word.apply_theme",      self.apply_theme,      "Apply a document theme", ["theme_name"])
        self.register_tool("word.save_document",    self.save_document,    "Save the document", ["path"], risk_level=RiskLevel.MEDIUM)
        self.register_tool("word.close_document",   self.close_document,   "Close the active document")
        self.register_tool("word.insert_toc",       self.insert_toc,       "Insert table of contents")
        self.register_tool("word.insert_title_page",self.insert_title_page,"Insert a formatted title page", ["title", "subtitle"])

    # ─────────────────────────────────────────
    # Document Lifecycle
    # ─────────────────────────────────────────

    def create_document(self, template: str = None) -> Dict[str, Any]:
        """Create a new Word document."""
        if self._use_com:
            return self._create_com(template)
        elif _DOCX_AVAILABLE:
            return self._create_docx(template)
        raise RuntimeError("No Word library available")

    def _create_com(self, template: str = None) -> Dict[str, Any]:
        try:
            pythoncom.CoInitialize()
            self._word_app = win32.Dispatch("Word.Application")
            self._word_app.Visible = False
            if template and Path(template).exists():
                self._document = self._word_app.Documents.Add(Template=template)
            else:
                self._document = self._word_app.Documents.Add()
            logger.info("Created document via COM")
            return {"created": True, "method": "com"}
        except Exception as e:
            logger.warning(f"COM create failed ({e}), falling back to python-docx")
            self._word_app = None
            return self._create_docx(template)

    def _create_docx(self, template: str = None) -> Dict[str, Any]:
        if template and Path(template).exists():
            self._document = Document(template)
        else:
            self._document = Document()
            # Apply professional default styles
            self._apply_default_styles()
        logger.info("Created document via python-docx")
        return {"created": True, "method": "docx"}

    def _apply_default_styles(self):
        """Set up professional default styles."""
        if not _DOCX_AVAILABLE or self._word_app:
            return
        doc = self._document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def open_document(self, path: str) -> Dict[str, Any]:
        """Open an existing Word document."""
        resolved = str(Path(path).expanduser().resolve())
        if not Path(resolved).exists():
            raise FileNotFoundError(f"Word file not found: {resolved}")
        self._doc_path = resolved

        if self._use_com:
            try:
                pythoncom.CoInitialize()
                self._word_app = win32.Dispatch("Word.Application")
                self._word_app.Visible = False
                self._document = self._word_app.Documents.Open(resolved)
                return {"opened": True, "path": resolved, "method": "com"}
            except Exception as e:
                logger.warning(f"COM open failed: {e}")

        if _DOCX_AVAILABLE:
            self._document = Document(resolved)
            return {"opened": True, "path": resolved, "method": "docx"}

        raise RuntimeError("No Word library available")

    def close_document(self, save: bool = False) -> Dict[str, Any]:
        """Close the active document."""
        if self._word_app and self._document:
            try:
                self._document.Close(SaveChanges=-1 if save else 0)
                self._word_app.Quit()
            except Exception as e:
                logger.warning(f"COM close error: {e}")
            finally:
                self._word_app = None
                self._document = None
        elif self._document:
            self._document = None
        return {"closed": True}

    # ─────────────────────────────────────────
    # Content Insertion
    # ─────────────────────────────────────────

    def insert_title_page(
        self,
        title: str,
        subtitle: str = "",
        date_str: str = None,
        author: str = "",
    ) -> Dict[str, Any]:
        """Insert a formatted title page."""
        from datetime import date
        date_text = date_str or date.today().strftime("%B %d, %Y")

        if self._word_app:
            doc = self._document
            sel = self._word_app.Selection
            # Title
            sel.Style = doc.Styles("Title")
            sel.TypeText(title)
            sel.TypeParagraph()
            if subtitle:
                sel.Style = doc.Styles("Subtitle")
                sel.TypeText(subtitle)
                sel.TypeParagraph()
            sel.TypeText(date_text)
            sel.TypeParagraph()
            if author:
                sel.TypeText(author)
                sel.TypeParagraph()
            # Page break
            sel.InsertBreak(7)
        else:
            doc = self._document
            p = doc.add_heading(title, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if subtitle:
                sp = doc.add_paragraph(subtitle)
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp.runs[0].font.size = Pt(16)
                sp.runs[0].bold = True
            dp = doc.add_paragraph(date_text)
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if author:
                ap = doc.add_paragraph(author)
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

        return {"inserted": "title_page", "title": title}

    def insert_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """Insert a heading at the specified level (1-9)."""
        level = max(1, min(9, level))
        if self._word_app:
            sel = self._word_app.Selection
            sel.Style = self._document.Styles(f"Heading {level}")
            sel.TypeText(text)
            sel.TypeParagraph()
        else:
            self._document.add_heading(text, level=level)
        return {"inserted": "heading", "text": text, "level": level}

    def insert_paragraph(
        self,
        text: str,
        style: str = "Normal",
        bold: bool = False,
        italic: bool = False,
        font_size: int = None,
    ) -> Dict[str, Any]:
        """Insert a paragraph with optional formatting."""
        if self._word_app:
            sel = self._word_app.Selection
            try:
                sel.Style = self._document.Styles(style)
            except Exception:
                pass
            if bold:
                sel.Font.Bold = True
            if italic:
                sel.Font.Italic = True
            if font_size:
                sel.Font.Size = font_size
            sel.TypeText(text)
            sel.TypeParagraph()
            if bold:
                sel.Font.Bold = False
            if italic:
                sel.Font.Italic = False
        else:
            p = self._document.add_paragraph(text, style=style if style != "Normal" else None)
            if bold or italic or font_size:
                for run in p.runs:
                    if bold:
                        run.bold = True
                    if italic:
                        run.italic = True
                    if font_size:
                        run.font.size = Pt(font_size)

        return {"inserted": "paragraph", "length": len(text)}

    def insert_table(
        self,
        data: Any,
        headers: Any = None,
        style: str = "Table Grid",
    ) -> Dict[str, Any]:
        """Insert a formatted table from list-of-dicts data.

        `data` should be a list of dicts, but this method defensively handles:
        - A Python list of dicts (normal)
        - A single dict (wraps in list)
        - A stringified list/dict (e.g. from imperfect template resolution)
        """
        # ── Normalize data ──────────────────────────────────────────────
        if isinstance(data, str):
            # Try to parse a stringified Python literal or JSON
            import ast, json
            cleaned = data.strip()
            # Wrap bare dict/list of dicts in brackets if needed
            if cleaned.startswith("{"):
                cleaned = f"[{cleaned}]"
            try:
                data = ast.literal_eval(cleaned)
            except Exception:
                try:
                    data = json.loads(cleaned)
                except Exception:
                    logger.warning(f"insert_table: could not parse data string, got: {cleaned[:200]}")
                    return {"inserted": "table", "rows": 0, "error": "Could not parse data"}

        if isinstance(data, dict):
            data = [data]

        if not data or not isinstance(data, list):
            # Raise instead of silently returning — makes failure visible in logs
            # so the orchestrator can re-plan instead of saving an empty doc.
            raise ValueError(
                f"insert_table received empty or invalid data: {repr(data)[:200]}. "
                "Check that the template variable resolves to a list of dicts."
            )

        # ── Normalize headers ───────────────────────────────────────────
        if isinstance(headers, str):
            # Could be "Payment_Mode, sum_Total_Amount" or "['a', 'b']"
            import ast
            h = headers.strip()
            if h.startswith("["):
                try:
                    headers = ast.literal_eval(h)
                except Exception:
                    headers = [x.strip().strip("'\"") for x in h.strip("[]").split(",")]
            else:
                headers = [x.strip() for x in h.split(",")]

        if headers is None:
            headers = list(data[0].keys()) if data else []

        # Ensure all rows are dicts
        if data and not isinstance(data[0], dict):
            # Might be list of lists — convert using headers
            data = [dict(zip(headers, row)) for row in data if isinstance(row, (list, tuple))]

        rows_data = [[str(row.get(h, "")) for h in headers] for row in data]

        if self._word_app:
            try:
                self._insert_table_com(headers, rows_data, style)
            except Exception as e:
                logger.warning(f"COM table insert failed ({e}), falling back to python-docx path")
                # Save COM doc to disk, reopen with python-docx, insert table, save back
                if self._doc_path and _DOCX_AVAILABLE:
                    self._table_via_docx_on_saved_file(headers, rows_data, style)
                else:
                    raise
        else:
            self._insert_table_docx(headers, rows_data, style)

        return {
            "inserted": "table",
            "headers": headers,
            "rows": len(rows_data),
            "columns": len(headers),
        }

    def _insert_table_com(self, headers, rows, style):
        doc = self._document
        # Move cursor to the very end of the document so the table is always
        # appended after existing content, not inserted at a random position.
        wdStory = 6  # Word constant: move to end of document
        self._word_app.Selection.EndKey(Unit=wdStory)

        sel = self._word_app.Selection
        total_rows = len(rows) + 1
        total_cols = len(headers)
        table = doc.Tables.Add(sel.Range, total_rows, total_cols)
        try:
            table.Style = style
        except Exception:
            pass
        try:
            table.ApplyStyleHeadingRows = True
        except Exception:
            pass

        for c, h in enumerate(headers):
            cell = table.Cell(1, c + 1)
            cell.Range.Text = str(h)
            cell.Range.Font.Bold = True

        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                table.Cell(r + 2, c + 1).Range.Text = str(val)

    def _table_via_docx_on_saved_file(self, headers, rows, style):
        """Fallback: save COM doc, add table via python-docx, re-save."""
        import tempfile, shutil
        tmp = self._doc_path + ".tmp.docx"
        try:
            self._document.SaveAs2(tmp, FileFormat=16)
            doc = Document(tmp)
            # Append table
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            try:
                table.style = style
            except Exception:
                pass
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = str(h)
                if hdr_cells[i].paragraphs[0].runs:
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
            for r_idx, row in enumerate(rows):
                cells = table.rows[r_idx + 1].cells
                for c_idx, val in enumerate(row):
                    cells[c_idx].text = str(val)
            doc.save(tmp)
            shutil.copy2(tmp, self._doc_path)
            logger.info(f"Table inserted via python-docx fallback into {self._doc_path}")
        finally:
            try:
                import os
                os.remove(tmp)
            except Exception:
                pass

    def _insert_table_docx(self, headers, rows, style):
        doc = self._document
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        try:
            table.style = style
        except Exception:
            pass

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            run = hdr_cells[i].paragraphs[0].runs
            if run:
                run[0].bold = True
                run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row_data):
                cells[c_idx].text = val

    def insert_image(
        self,
        image_path: str,
        caption: str = None,
        width_inches: float = 5.0,
    ) -> Dict[str, Any]:
        """Insert an image into the document."""
        resolved = str(Path(image_path).expanduser().resolve())
        if not Path(resolved).exists():
            return {"inserted": False, "error": f"Image not found: {resolved}"}

        if self._word_app:
            try:
                rng = self._word_app.Selection.Range
                pic = self._document.InlineShapes.AddPicture(
                    FileName=resolved, LinkToFile=False,
                    SaveWithDocument=True, Range=rng
                )
                pic.Width = width_inches * 72
                pic.Height = pic.Width * (pic.Height / pic.Width)
            except Exception as e:
                logger.warning(f"COM image insert failed: {e}")
                return {"inserted": False, "error": str(e)}
        else:
            self._document.add_picture(resolved, width=Inches(width_inches))
            if caption:
                cap = self._document.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True

        return {"inserted": "image", "path": resolved}

    def insert_page_break(self) -> Dict[str, Any]:
        """Insert a page break."""
        if self._word_app:
            self._word_app.Selection.InsertBreak(7)
        else:
            self._document.add_page_break()
        return {"inserted": "page_break"}

    def insert_toc(self) -> Dict[str, Any]:
        """Insert an automatic table of contents."""
        if self._word_app:
            try:
                rng = self._word_app.Selection.Range
                self._document.TablesOfContents.Add(
                    Range=rng, UseHeadingStyles=True,
                    UpperHeadingLevel=1, LowerHeadingLevel=3,
                    IncludePageNumbers=True,
                )
                return {"inserted": "toc"}
            except Exception as e:
                return {"inserted": False, "error": str(e)}
        else:
            # python-docx doesn't support TOC natively — insert placeholder
            p = self._document.add_paragraph()
            run = p.add_run("[Table of Contents]")
            run.bold = True
            return {"inserted": "toc_placeholder"}

    def set_header(self, text: str) -> Dict[str, Any]:
        """Set the document header."""
        if self._word_app:
            section = self._document.Sections(1)
            header = section.Headers(1)
            header.Range.Text = text
        else:
            section = self._document.sections[0]
            header = section.header
            header.paragraphs[0].text = text
        return {"set": "header", "text": text}

    def set_footer(self, text: str) -> Dict[str, Any]:
        """Set the document footer."""
        if self._word_app:
            section = self._document.Sections(1)
            footer = section.Footers(1)
            footer.Range.Text = text
        else:
            section = self._document.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = text
        return {"set": "footer", "text": text}

    def apply_theme(self, theme_name: str = "Office") -> Dict[str, Any]:
        """Apply a document theme (COM only)."""
        if not self._word_app:
            return {"applied": False, "reason": "COM required for themes"}
        try:
            self._document.ApplyTheme(theme_name)
            return {"applied": True, "theme": theme_name}
        except Exception as e:
            return {"applied": False, "error": str(e)}

    # ─────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────

    def save_document(self, path: str, format: str = "docx") -> Dict[str, Any]:
        """Save the document to disk."""
        resolved = str(Path(path).expanduser().resolve())
        ensure_dir(Path(resolved).parent)
        self._doc_path = resolved

        if self._word_app:
            fmt_map = {"docx": 16, "pdf": 17, "txt": 2, "rtf": 6}
            fmt_code = fmt_map.get(format.lower(), 16)
            self._document.SaveAs2(resolved, FileFormat=fmt_code)
        else:
            if format.lower() == "pdf":
                logger.warning("PDF export requires COM. Saving as DOCX instead.")
                resolved = resolved.replace(".pdf", ".docx")
            self._document.save(resolved)

        size = Path(resolved).stat().st_size if Path(resolved).exists() else 0
        return {
            "saved": True,
            "path": resolved,
            "format": format,
            "size_bytes": size,
        }

    def __del__(self):
        try:
            if self._word_app:
                self._document.Close(SaveChanges=0)
                self._word_app.Quit()
        except Exception:
            pass
